"""Build CLIM715_Experiment_Cheatsheet.docx - a today-ready handout
explaining exactly how the CLIM-715 ground-flux experiment was set up
and run. For explaining the project to others.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PURPLE = RGBColor(0x6B, 0x5C, 0xB5)
TEAL = RGBColor(0x2C, 0x8A, 0x83)
ORANGE = RGBColor(0xC2, 0x49, 0x2A)
GREY = RGBColor(0x4A, 0x50, 0x66)
DARK = RGBColor(0x2C, 0x31, 0x42)

doc = Document()

# Page setup
for s in doc.sections:
    s.left_margin = Inches(0.85)
    s.right_margin = Inches(0.85)
    s.top_margin = Inches(0.7)
    s.bottom_margin = Inches(0.7)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_title(text, color=PURPLE, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def add_subtitle(text, color=GREY, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def add_h(text, level=1, color=PURPLE):
    sizes = {1: 16, 2: 13, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 11))
    r.font.color.rgb = color
    return p

def add_para(text, bold_first_phrase=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_first_phrase:
        r1 = p.add_run(bold_first_phrase + ' ')
        r1.bold = True
        r1.font.color.rgb = DARK
        r2 = p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_bullet(text, bold_lead=None, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r1 = p.add_run(bold_lead + ' ')
        r1.bold = True
        r1.font.color.rgb = DARK
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_num(text, bold_lead=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r1 = p.add_run(bold_lead + ' ')
        r1.bold = True
        r1.font.color.rgb = DARK
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_equation_block(text):
    """Centred italic monospaced equation line (poor man's math)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = 'Cambria Math'
    r.font.size = Pt(13)
    r.font.color.rgb = DARK
    return p

def add_code_block(lines):
    """Monospace block, indented, light background-ish via grey colour."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(10)
        r.font.color.rgb = DARK

def add_callout(text, color=TEAL):
    """Indented, italic, coloured callout."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = color
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.autofit = True

    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        c = hdr_cells[i]
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        # purple header background
        tc_pr = c._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '6B5CB5')
        tc_pr.append(shd)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = DARK

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = w
    return t

# =================================================================
# CONTENT BEGINS
# =================================================================

add_title("CLIM-715 Ground Flux Experiment", color=PURPLE, size=22)
add_subtitle("Everything you need to explain the experiment in one sitting.", size=12)
add_subtitle("Source: run_experiments.py + modified_full_report_1.docx. Built 2026-05-11.", size=9)

# -----------------------------------------------------------------
add_h("1. The Big Picture (30-second version)", level=1)
add_para(
    "We simulate one-dimensional heat conduction through three urban substrate columns "
    "(asphalt road, concrete roof, bare soil), each 2 metres deep, driven by a synthetic "
    "diurnal cycle of solar radiation and air temperature. We solve the heat equation using "
    "three different finite-difference schemes (FTCS, BTCS, Crank-Nicolson) at three "
    "different time steps (15 s, 60 s, 600 s). The surface temperature is computed at every "
    "step by solving the surface energy balance with Newton iteration. We then measure how "
    "the schemes' surface ground heat flux deviates from a fine-Δt reference.",
)
add_callout(
    "Headline finding: at Δt = 600 s, BTCS over-amplifies the diurnal G amplitude by 40 % on "
    "asphalt and concrete, 19 % on bare soil; Crank-Nicolson halves these errors. The "
    "dominant error is operator-splitting between the column and SEB updates."
)

# -----------------------------------------------------------------
add_h("2. Sequence of Steps — How the Experiment Runs", level=1)
add_para("The whole project follows this order; everything else is detail.")

add_num("Define three substrates as layered stacks of (depth, λ, C).", bold_lead="Define substrates.")
add_num("Build a stretched vertical grid for each substrate (top cell 0.5 cm for asphalt and concrete, 1 cm for bare soil; geometric stretch 1.18-1.25 down to bottom cell ≈ 30 cm at z = 2 m).", bold_lead="Build the grid.")
add_num("Initialize substrate temperatures from the analytical damping-depth profile.", bold_lead="Initialize.")
add_num("Loop over time steps. For each step:", bold_lead="Time-step loop.")
add_bullet("Sub-step A: advance the column from time n to n+1 using the θ-method, holding Tₛ⁰ fixed.", level=1)
add_bullet("Sub-step B: solve the surface energy balance for new Tₛ⁰ via Newton iteration, holding the column fixed.", level=1)
add_bullet("Apply bottom BC: T[N-1] = T[N-2] (zero-flux Neumann).", level=1)
add_num("Run for 2 diurnal cycles. Use day-2 (after 1-day spin-up) for all diagnostics.", bold_lead="Integration period.")
add_num("Compare each scheme's day-2 surface ground heat flux G(t) and surface temperature Tₛ⁰(t) against the reference run (FTCS at Δt = 15 s).", bold_lead="Compare.")
add_num("Report: RMSE of Tₛ⁰, peak-amplitude ratio of G, daily-storage ratio.", bold_lead="Diagnose.")

# -----------------------------------------------------------------
add_h("3. The Three Substrates", level=1)
add_para("Each substrate is a 2 m deep column of layered material with prescribed albedo at the top.")
add_table(
    ["Substrate", "Layers (depth, λ W/m/K, C ×10⁶ J/m³/K)", "Albedo α_s", "d (cm)"],
    [
        ["Asphalt road (4 layers)", "Asphalt 0-5 cm (0.75, 2.0); aggregate 5-25 cm (1.40, 2.4); dry soil 25-100 cm (0.30, 1.3); subsoil 100-200 cm (0.50, 1.8)", "0.10", "10.2"],
        ["Concrete roof (3 layers)", "Concrete deck 0-10 cm (1.50, 2.1); mineral-wool insulation 10-20 cm (0.04, 0.08); drywall/wood 20-200 cm (0.15, 1.5)", "0.30", "14.0"],
        ["Bare soil (uniform)", "Sandy loam 0-200 cm (0.30, 1.3)", "0.20", "8.0"],
    ],
)
add_callout("d = √(2κ/ω) is the diurnal damping depth, the e-folding depth of the daily wave. Computed from each substrate's top-layer κ = λ/C.")

# -----------------------------------------------------------------
add_h("4. The Governing Equation (the physics)", level=1)
add_para(
    "The heat equation in conductivity form (preserves discrete heat conservation across layer interfaces):",
)
add_equation_block("C_s(z) · ∂T_s/∂t  =  ∂/∂z [ λ_s(z) · ∂T_s/∂z ]")
add_para(
    "In words: the volumetric heat capacity times the rate of temperature change at depth z "
    "equals the spatial divergence of the conductive heat flux. Heat piles up wherever more "
    "flux enters from above than exits below.",
)

# -----------------------------------------------------------------
add_h("5. The Three Schemes — One Equation, Three α Values", level=1)
add_para("All three time-stepping schemes are special cases of a single θ-method update equation:")
add_equation_block("C_j Δz_j (T_j^(n+1) − T_j^n) / Δt = α [−∂G/∂z]^(n+1) + (1−α) [−∂G/∂z]^n")
add_para("where the flux-divergence at cell j is computed from G_{j±1/2} = λ_{j±1/2} (T_j − T_{j±1}) / (centre-to-centre spacing). Choose α; get a scheme:")
add_table(
    ["α", "Scheme", "Stability", "Order in Δt", "Per-step cost"],
    [
        ["0",   "FTCS  (Forward in Time)",  "Conditionally stable: ν ≤ 1/2",   "1st", "Cheap (vectorised arithmetic)"],
        ["1/2", "Crank-Nicolson",           "Unconditionally stable",          "2nd", "Tridiagonal solve, O(N)"],
        ["1",   "BTCS  (Backward in Time)", "Unconditionally stable, damping", "1st", "Tridiagonal solve, O(N)"],
    ],
)
add_callout("ν = κ Δt / Δz² is the dimensionless diffusion number — the diffusion analogue of the Courant-Friedrichs-Lewy number.")

# -----------------------------------------------------------------
add_h("6. Boundary Conditions — What Closes the System", level=1)
add_para(
    "The heat equation is second-order in z, so we need two BCs (one at each end of the column).",
)

add_h("6a. Top BC — Surface Energy Balance (SEB) closure", level=2)
add_para("The top of the column is closed by the requirement that the four surface fluxes balance to zero:")
add_equation_block("R_n − H − LE − G = 0   (solved for T_s⁰)")
add_para("with each flux evaluated at the surface temperature T_s⁰:")
add_bullet("R_n = (1 − α_s) S↓ + ε_s L↓ − ε_s σ (T_s⁰)⁴   (net radiation; σ = 5.67×10⁻⁸ W/m²/K⁴)", bold_lead="Net radiation.")
add_bullet("H = ρ c_p (T_s⁰ − T_a) / r_a    where r_a = 1 / (C_H U) ≈ 67 s/m   (sensible heat flux)", bold_lead="Sensible.")
add_bullet("LE = 0   (strict-impervious assumption — no evaporation)", bold_lead="Latent.")
add_bullet("G = λ_{1/2} (T_s⁰ − T_1) / (z_1 − z_0)   (ground heat flux, from the top half-level)", bold_lead="Ground.")
add_para(
    "The (T_s⁰)⁴ term makes this equation nonlinear in T_s⁰. We solve it by Newton iteration:",
)
add_equation_block("T_s⁰_new = T_s⁰_old − F(T_s⁰_old) / F′(T_s⁰_old)")
add_para("where F = R_n − H − LE − G. Typically converges in 3-4 iterations to |ΔT| < 10⁻⁴ K.")

add_h("6b. Bottom BC — Zero-flux Neumann at z = 2 m", level=2)
add_para("At the deepest face, the temperature gradient is set to zero:")
add_equation_block("∂T_s/∂z |_{z=2m} = 0    ⟺    G|_{z=2m} = 0    ⟺    T_{N-1} = T_{N-2}")
add_para(
    "Discretely, one line of code: the deepest cell mirrors the one above it. Justified "
    "because z = 2 m is 14-25 damping depths down, where the diurnal amplitude is already "
    "negligible (~10⁻⁶ of the surface value).",
)

# -----------------------------------------------------------------
add_h("7. Von Neumann Stability Analysis — How We Used It", level=1)
add_para("We used von Neumann analysis to PREDICT which (scheme, Δt, substrate) combinations would blow up, then VERIFIED the prediction empirically in Test 1 and Test 2.")
add_h("7a. The procedure (four steps)", level=2)
add_num("Substitute a Fourier mode  T_j^n = A^n · e^(ikjΔz)  into each scheme's update equation.")
add_num("Cancel common factors and solve for the per-step amplification factor A as a function of ν and kΔz.")
add_num("Demand |A| ≤ 1 for every resolvable wavenumber, i.e., kΔz ∈ [0, π].")
add_num("Note that the worst case is always at kΔz = π (the 2-Δz wave, where adjacent cells flip sign). Substitute that and solve for the stability bound on ν.")

add_h("7b. Results — the three amplification factors at kΔz = π", level=2)
add_table(
    ["Scheme", "A at worst wave", "|A| over ν > 0", "Stability bound"],
    [
        ["FTCS",            "1 − 4ν",        "> 1 when ν > 1/2",   "Conditional: ν ≤ 1/2"],
        ["BTCS",            "1/(1 + 4ν)",    "Always in (0, 1]",   "Unconditional"],
        ["Crank-Nicolson",  "(1 − 2ν)/(1 + 2ν)", "Always in [−1, 1]", "Unconditional"],
    ],
)

add_h("7c. The 27-cell FTCS prediction (Test 2)", level=2)
add_para("Plugging each substrate's top-cell κ, top-cell Δz, and the three Δt values into ν = κΔt/Δz² gives the FTCS stability picture in advance:")
add_table(
    ["Substrate (top κ, Δz₀)", "ν at Δt=15s", "ν at Δt=60s", "ν at Δt=600s"],
    [
        ["Bare soil  (2.31e-7 m²/s, 1.0 cm)",     "0.035 ✓",       "0.139 ✓",      "1.39 ✗ (mild)"],
        ["Asphalt road (3.75e-7 m²/s, 0.5 cm)",   "0.225 ✓",       "0.900 ✗",      "9.00 ✗ (catastrophic)"],
        ["Concrete roof (7.14e-7 m²/s, 0.5 cm)",  "0.428 ⚠",      "1.71 ✗",       "17.14 ✗ (catastrophic)"],
    ],
)
add_callout("Test 2 then confirmed these predictions: FTCS blew up at Δt=60s on asphalt and concrete, at Δt=600s on all three substrates (mildly on soil). BTCS and CN completed every run.")

# -----------------------------------------------------------------
add_h("8. The Two Experiments — Test 1 and Test 2", level=1)

add_h("Test 1 — Verification (does the solver work?)", level=2)
add_table(
    ["Setting", "Value"],
    [
        ["Substrate",          "Uniform sandy loam (1 material, no layers)"],
        ["Grid",               "Uniform Δz = 1 cm everywhere"],
        ["Surface BC",         "Prescribed: T_s⁰(t) = 292.5 + 7.5·cos(ω t) K"],
        ["Bottom BC",          "Zero-flux Neumann at z = 2 m"],
        ["Reference",          "Closed-form analytical damping-depth solution"],
        ["Integration",        "5 diurnal cycles; day-5 used for diagnostics"],
        ["Configurations",     "6: FTCS at ν=0.4 and ν=0.6; BTCS and CN at Δt=300 s and 900 s"],
        ["Question answered",  "Does the solver reproduce the textbook analytical answer? (Yes.)"],
    ],
)
add_para(
    "FTCS at ν = 0.6 was predicted to blow up. It did, at step 53 — matching the predicted "
    "geometric growth factor |1 − 4·0.6|^53 ≈ 5.5×10⁷ per step. Confirmation of the theory.",
)

add_h("Test 2 — Prognostic SEB (what we actually wanted to know)", level=2)
add_table(
    ["Setting", "Value"],
    [
        ["Substrate",          "Three: asphalt road, concrete roof, bare soil"],
        ["Grid",               "Stretched per-substrate: top 0.5-1 cm, geometric stretch 1.18-1.25, bottom ≈ 30 cm"],
        ["Surface BC",         "Prognostic: T_s⁰ from SEB Newton iteration (top BC §6a above)"],
        ["Bottom BC",          "Zero-flux Neumann at z = 2 m"],
        ["Reference",          "FTCS at Δt = 15 s on the same substrate"],
        ["Integration",        "2 diurnal cycles; day-2 used for diagnostics"],
        ["Configurations",     "27: 3 substrates × 3 schemes (FTCS/BTCS/CN) × 3 Δt (15 / 60 / 600 s)"],
        ["Question answered",  "How do FTCS, BTCS, CN compare at operational time steps on real urban substrates?"],
    ],
)

# -----------------------------------------------------------------
add_h("9. Algorithm — What Runs Inside One Time Step", level=1)
add_para("For each substrate × scheme α × Δt cell of the matrix, the following loop executes:")
add_code_block([
    "for step in range(N_steps):",
    "    # SUB-STEP A: column update (theta-method)",
    "    G_half = harmonic_mean_lambda * (T - T_neighbour) / dz_centre   # half-level fluxes",
    "    if alpha == 0:",
    "        # FTCS — explicit, no linear system",
    "        T_new = T + dt / (C * dz) * (G_half[:-1] - G_half[1:])",
    "    else:",
    "        # BTCS (alpha=1) or CN (alpha=0.5) — implicit, tridiagonal solve",
    "        A_band, b = build_tridiagonal_system(T, alpha, dt, dz, lambda_half, C)",
    "        T_new = scipy.linalg.solve_banded((1, 1), A_band, b)",
    "",
    "    # Apply bottom BC (zero-flux Neumann)",
    "    T_new[N-1] = T_new[N-2]",
    "",
    "    # SUB-STEP B: SEB update (Newton iteration for new T_s^0)",
    "    Ts0 = Ts0_old",
    "    for _ in range(MAX_ITER):",
    "        F  = R_n(Ts0) - H(Ts0) - LE - G(Ts0, T_new[1])",
    "        dF = dR_n/dT - dH/dT - dG/dT",
    "        dTs = -F / dF",
    "        Ts0 = Ts0 + dTs",
    "        if abs(dTs) < 1e-4: break",
    "",
    "    # Update top cell to the new surface temperature",
    "    T_new[0] = Ts0",
    "",
    "    # Advance time",
    "    T = T_new",
    "    t = t + dt",
    "    record(Ts0, G_half[0])",
])
add_callout(
    "Order of operations is the key design choice: column FIRST (holding T_s⁰ fixed), then "
    "SEB SECOND (holding column fixed). This is the operator splitting that produces the "
    "first-order error analysed in §5.2 of the report."
)

# -----------------------------------------------------------------
add_h("10. What the Results Showed (one-page summary)", level=1)

add_h("From Test 1 (verification)", level=2)
add_bullet("FTCS at ν = 0.4: RMSE_T = 0.007 K — essentially exact.")
add_bullet("FTCS at ν = 0.6: BLEW UP at step 53 — exactly as von Neumann predicted.")
add_bullet("BTCS error grows linearly with Δt (first-order). CN error stays flat (second-order behaviour at this regime).")
add_bullet("Conclusion: solver passes verification. Trust it on Test 2.")

add_h("From Test 2 (prognostic SEB)", level=2)
add_bullet("At Δt = 15 s: all three schemes agree to line-thickness on every substrate. Control condition.")
add_bullet("At Δt = 60 s: FTCS blows up on asphalt (ν=0.90) and concrete (ν=1.71). BTCS and CN run to completion.")
add_bullet("At Δt = 600 s: FTCS catastrophically unstable on asphalt and concrete; mildly unstable on soil. BTCS over-amplifies the diurnal G amplitude by 40 % / 41 % / 19 % (asphalt / roof / soil). CN halves these errors to 21 % / 26 % / 10 %.")
add_bullet("Δt-refinement ratios: all six BTCS and CN values cluster in [8.2, 10.2] — first-order in Δt for both schemes, even though CN is intrinsically second-order. This is the smoking-gun signature of operator-splitting error.")
add_bullet("SHAP attribution on a 150-column synthetic ensemble: κ_top is the dominant predictor of BTCS coarse-Δt error, with R² = 0.92 from κ_top alone.")
add_bullet("Daily storage integral preserved within ±5 % across all schemes — so the over-amplification is symmetric (daytime overshoot + nighttime undershoot), preserving the daily mean. Implication: UHI diurnal range is biased, but UHI daily mean is preserved.")

add_h("Why the operator-splitting error dominates", level=2)
add_para(
    "At every step, sub-step A advances the column with the OLD surface temperature, and "
    "sub-step B then updates the surface temperature with the NEW column. The column "
    "therefore equilibrates to the wrong T_s⁰ during sub-step A, and the discrepancy "
    "manifests as a spurious top-cell gradient. This error is first-order in Δt regardless "
    "of the per-substep scheme — which is why CN (intrinsically second-order) shows the same "
    "ratio of 10 as BTCS (first-order) when Δt scales by 10."
)
add_para(
    "Two fixes were sketched in §5.4 of the report but not implemented: a fully coupled SEB "
    "row in the tridiagonal solve, or a Strang-symmetric splitting (A → B → A with half-step "
    "intervals). Either would eliminate the leading first-order splitting error.",
)

# =================================================================
# TEST 1 DEEP DIVE - addressing follow-up confusion
# =================================================================

doc.add_page_break()
add_title("Part II — Test 1 Deep Dive", color=PURPLE, size=20)
add_subtitle("Walking through Test 1 step by step. Every number, every line of the algorithm, every boundary condition application.")

# -----------------------------------------------------------------
add_h("11. Test 1 — The Setup, in One Place", level=1)
add_para("Test 1 is the simplest possible configuration. Use it to understand the solver before tackling Test 2.")
add_table(
    ["Setting", "Value", "Why"],
    [
        ["Substrate",          "Uniform sandy loam (single material)",                  "No layer jumps to worry about"],
        ["λ (conductivity)",   "0.30 W/m/K",                                            "Standard sandy-loam value"],
        ["C (heat capacity)",  "1.3 × 10⁶ J/m³/K",                                      "Standard sandy-loam value"],
        ["κ = λ/C (diffusivity)", "2.31 × 10⁻⁷ m²/s",                                    "Derived from λ, C"],
        ["d = √(2κ/ω) (damping depth)", "0.0797 m = 7.97 cm",                            "Daily wave decays to 37 % at this depth"],
        ["Grid",               "Uniform Δz = 1 cm, 201 cells (z = 0, 1, 2, ..., 200 cm)", "Clean comparison to analytical solution"],
        ["T_mean",             "292.5 K (= 19.35 °C)",                                  "Daily-mean surface temperature"],
        ["A_0",                "7.5 K",                                                 "Half the peak-to-trough surface swing"],
        ["ω",                  "2π/86400 ≈ 7.27 × 10⁻⁵ rad/s",                          "One full diurnal cycle per day"],
        ["Top BC",             "T_s⁰(t) = 292.5 + 7.5 · cos(ω t) K  [PRESCRIBED]",      "Set explicitly; no SEB needed"],
        ["Bottom BC",          "T[N-1] = T[N-2]  [zero-flux Neumann at z = 2 m]",       "Damped signal already negligible"],
        ["Initial condition",  "Analytical damping-depth profile at t = 0",             "Skip spin-up, start in quasi-equilibrium"],
        ["Integration",        "5 days; day-5 used for diagnostics",                    "Any transient from non-perfect IC has decayed"],
    ],
)

# -----------------------------------------------------------------
add_h("12. Where Does the Initial Temperature at Each Cell Come From?", level=1)
add_para("The initial condition is used ONCE at t = 0. After that, the IC has no further role — the simulation is driven forward by the BCs and the θ-method update.")

add_h("12a. The formula", level=2)
add_para("We set every cell's temperature using the analytical damping-depth solution evaluated at t = 0:")
add_equation_block("T(z, 0)  =  T_mean  +  A₀ · exp(−z / d) · cos(−z / d)")
add_para("Two pieces multiply A₀: the exponential envelope exp(−z/d) which shrinks amplitude with depth, and the cosine cos(−z/d) which encodes the phase delay with depth.")

add_h("12b. Worked-out initial values for Test 1", level=2)
add_para("Plugging in T_mean = 292.5 K, A_0 = 7.5 K, d = 7.97 cm:")
add_table(
    ["Cell", "Depth z (cm)", "z/d", "exp(−z/d)", "cos(−z/d)", "T(z, 0) (K)"],
    [
        ["Cell 0 (surface)", "0",   "0.00",  "1.000", "1.000",  "300.00"],
        ["Cell 1",           "1",   "0.125", "0.882", "0.992",  "299.07"],
        ["Cell 5",           "5",   "0.627", "0.535", "0.810",  "295.75"],
        ["Cell 8",           "8",   "1.00",  "0.368", "0.540",  "294.00"],
        ["Cell 10 (≈ d)",    "10",  "1.255", "0.286", "0.310",  "293.16"],
        ["Cell 20",          "20",  "2.510", "0.081", "−0.808", "292.01"],
        ["Cell 50",          "50",  "6.276", "0.0019", "0.999", "292.51"],
        ["Cell 100",         "100", "12.55",  "3.5e-6", "0.97", "292.50"],
        ["Cell 200 (bottom)", "200", "25.10", "1.2e-11", "0.99", "292.50"],
    ],
)
add_callout(
    "By cell 50 (50 cm down — about 6 damping depths), the initial T is already indistinguishable "
    "from the daily mean 292.5 K. By cell 200 (2 m down), it is 292.500000... — flat to machine "
    "precision. That is why the column is 2 m deep: 14-25 damping depths is overkill, but it gives "
    "the bottom BC nothing physical to do."
)

add_h("12c. Why this IC and not just zero (or T = T_mean) everywhere?", level=2)
add_para(
    "If we initialised every cell at T = T_mean uniformly, the column would not be in equilibrium "
    "with the diurnal forcing. The first few simulation days would be spent letting the daily wave "
    "penetrate from the surface and the column adjust. That spin-up phase is wasted compute. By "
    "starting from the analytical t = 0 profile, the column is already in quasi-equilibrium and "
    "day-2 (or day-5) diagnostics are clean."
)
add_para(
    "Note: \"quasi-equilibrium\" — not exact — because our numerical scheme is not the exact "
    "analytical solution. There is a tiny mismatch between the IC and what the numerical solver "
    "would have produced at t = 0 if it had been running forever. That mismatch decays within a "
    "few diurnal cycles, which is why we wait until day 5 (Test 1) or day 2 (Test 2) for "
    "diagnostics."
)

# -----------------------------------------------------------------
add_h("13. What is harmonic_mean_lambda?", level=1)
add_para(
    "harmonic_mean_lambda is the effective thermal conductivity at a FACE between two adjacent cells. "
    "It is computed from the two cell-centre conductivities by the harmonic-mean formula:"
)
add_equation_block("λ_{j+½}  =  2 · λ_j · λ_{j+1} / (λ_j + λ_{j+1})")

add_h("13a. Why harmonic mean (not arithmetic mean)?", level=2)
add_para(
    "Derivation: consider two materials in series, each obeying Fourier's law. In steady state, the "
    "SAME heat flux must flow through both. The temperature drop divides between them in inverse "
    "proportion to their conductivities. Solving gives the harmonic-mean formula, which guarantees "
    "the discrete flux at the face exactly matches the continuous flux."
)
add_para("Numerical example — concrete deck (λ = 1.50) on mineral-wool insulation (λ = 0.04):")
add_bullet("Arithmetic mean: (1.50 + 0.04) / 2 = 0.77 — over-estimates the flux by a factor of 10.",
           bold_lead="Arithmetic.")
add_bullet("Harmonic mean: 2 · 1.50 · 0.04 / (1.50 + 0.04) = 0.078 — correctly captures the insulator's choke-point behaviour.",
           bold_lead="Harmonic.")
add_callout(
    "The harmonic mean ensures the discrete code preserves heat conservation across material "
    "interfaces. Arithmetic mean would silently destroy energy conservation at every interface "
    "between materials with different λ."
)

add_h("13b. What does harmonic_mean_lambda equal for Test 1?", level=2)
add_para(
    "Test 1 uses a single uniform material (sandy loam with λ = 0.30 everywhere). At every face j+½:"
)
add_equation_block("λ_{j+½}  =  2 · 0.30 · 0.30 / (0.30 + 0.30)  =  0.18 / 0.60  =  0.30 W/m/K")
add_para(
    "So for Test 1, the harmonic mean is trivial — every face has λ = 0.30, the same value as every "
    "cell centre. The harmonic-mean formula does not do anything interesting here. But we still call "
    "the same function in the code, because the SAME solver runs for both Test 1 and Test 2; we do "
    "not have one solver for uniform substrates and a different one for layered substrates."
)
add_para("For Test 2 (layered substrates), the harmonic mean matters at every layer interface.")

# -----------------------------------------------------------------
add_h("14. Test 1 — The Algorithm at Every Time Step", level=1)
add_para(
    "Below is exactly what runs at every time step in Test 1, in order. Read it as a recipe."
)

add_h("Step 1: Compute the new prescribed surface temperature", level=2)
add_para("Advance time, then plug into the prescribed BC formula:")
add_code_block([
    "t_new = t_old + dt",
    "T_s_0_new = T_mean + A_0 * cos(omega * t_new)",
])
add_para(
    "This is NOT solved by Newton iteration in Test 1 — it is just plugged in. The surface "
    "temperature for Test 1 is externally prescribed, not solved from the SEB."
)

add_h("Step 2: Compute the harmonic-mean λ at every face", level=2)
add_code_block([
    "for j in range(N - 1):",
    "    lam_half[j] = 2 * lam[j] * lam[j+1] / (lam[j] + lam[j+1])",
    "# For Test 1: every lam_half[j] = 0.30 (uniform).",
    "# For Test 2: lam_half jumps at layer interfaces.",
])
add_para(
    "In practice, lam_half is precomputed ONCE at the start (it depends only on the substrate "
    "material distribution, not on T), so this step does not actually re-run every time."
)

add_h("Step 3: Update interior cells (j = 1, 2, ..., N−2) using the θ-method", level=2)
add_para("The θ-method update equation:")
add_equation_block("C · Δz · (T_j^(n+1) − T_j^n) / Δt  =  α · Q_j^(n+1)  +  (1−α) · Q_j^n")
add_para("where Q_j is the net flux divergence at cell j:")
add_equation_block("Q_j  =  ( λ_{j−½} · (T_{j−1} − T_j) / Δz )  −  ( λ_{j+½} · (T_j − T_{j+1}) / Δz )")

add_para("Two cases:")
add_bullet(
    "the right-hand side has only OLD values (time level n). We just compute:\n"
    "  T_new[j] = T[j] + (Δt / (C·Δz)) · ( λ_half[j-1]·(T[j-1]−T[j])/Δz − λ_half[j]·(T[j]−T[j+1])/Δz )\n"
    "One pass through the array — cheap.",
    bold_lead="α = 0 (FTCS):"
)
add_bullet(
    "the right-hand side has UNKNOWN T_new values too (time level n+1). We must solve a tridiagonal linear system:\n"
    "  A · T_new = b\n"
    "where A is built from −α, 1 + 2α (something), and the bottom and top BCs, and b is built from T (old).\n"
    "SciPy: T_new = scipy.linalg.solve_banded((1,1), A, b). O(N) cost.",
    bold_lead="α = 1 (BTCS) or α = 1/2 (CN):"
)

add_h("Step 4: Apply the bottom BC at every step", level=2)
add_code_block([
    "T_new[N-1] = T_new[N-2]   # zero-flux Neumann",
])
add_para(
    "Forces the gradient between the deepest two cells to zero → flux at the bottom face is zero. "
    "Reapplied at every step (not just once). Without this, the deep cells would drift."
)

add_h("Step 5: Apply the top BC at every step", level=2)
add_code_block([
    "T_new[0] = T_s_0_new   # prescribed surface value from Step 1",
])
add_para(
    "Plug in the value computed in Step 1. Reapplied at every step (the prescribed surface "
    "temperature changes sinusoidally throughout the day). Without this, the top cell would freely "
    "drift and the simulation would not represent the diurnal forcing."
)

add_h("Step 6: Advance and record", level=2)
add_code_block([
    "T = T_new",
    "t = t_new",
    "G_surface = lam_half[0] * (T[0] - T[1]) / dz   # surface ground heat flux",
    "record(t, T[0], G_surface, ...)",
    "if t < t_end: goto Step 1",
])

# -----------------------------------------------------------------
add_h("15. When Each Constraint Is Applied — Summary Table", level=1)
add_para(
    "The biggest source of confusion is whether the BCs are \"initial\" (one-time) or \"applied at every "
    "step\". Answer: BCs are applied at every step. The IC is applied only at t = 0."
)
add_table(
    ["Constraint", "When applied?", "What it does"],
    [
        ["Initial condition (IC)",       "ONCE, at t = 0 only",          "Sets every cell's starting T from the analytical damping-depth profile"],
        ["Top BC (Dirichlet, prescribed)", "AT EVERY TIME STEP",          "Sets T[0] to T_s⁰(t_new) = 292.5 + 7.5·cos(ω·t_new)"],
        ["Bottom BC (zero-flux Neumann)",  "AT EVERY TIME STEP",          "Sets T[N-1] = T[N-2], forcing zero flux at the bottom face"],
        ["Interior θ-method update",       "AT EVERY TIME STEP",          "Advances cells j = 1..N-2 from time n to n+1"],
        ["Harmonic-mean λ at faces",       "ONCE, at the start (precomputed)", "Used at every step inside the flux computation; depends only on the material map"],
    ],
)

# -----------------------------------------------------------------
add_h("16. Test 1 — Concrete Numbers, Step by Step (one Δt = 60 s step)", level=1)
add_para(
    "To make this absolutely concrete, here is what one Δt = 60 s step on Test 1 actually computes, "
    "starting from the analytical IC at t = 0:"
)

add_h("Before the first step (t = 0)", level=2)
add_bullet("T[0] = 300.00 K (set by the IC: T_mean + A_0·exp(0)·cos(0))")
add_bullet("T[1] = 299.07 K, T[5] = 295.75 K, T[10] = 293.16 K, T[200] = 292.50 K (also IC)")
add_bullet("All cells have λ = 0.30, so all λ_half = 0.30 (computed once at start)")

add_h("Step 1 — Compute new T_s⁰", level=2)
add_para("t_new = 60 s. ω · t_new = (7.27e-5)(60) = 0.00436 rad. cos(0.00436) ≈ 0.99999.")
add_bullet("T_s_0_new = 292.5 + 7.5 · 0.99999 ≈ 299.9999 K")
add_bullet("So the surface temperature barely changes in 60 s (still very near its peak).")

add_h("Step 2 — Harmonic-mean λ", level=2)
add_para("Already done: all λ_half = 0.30. No work this step.")

add_h("Step 3 — Update interior cells (using FTCS, α = 0)", level=2)
add_para("For ν = 0.4 (the safe FTCS case in Test 1), ν = κΔt/Δz² = (2.31e-7)(60)/(0.01)² = 0.139. Take Δt = 60 s. Per cell:")
add_para("T_new[j] = T[j] + ν · (T[j-1] − 2·T[j] + T[j+1])")
add_para("Worked example for cell j = 5:")
add_para("T_new[5] = 295.75 + 0.139 · (296.50 − 2·295.75 + 294.78) = 295.75 + 0.139 · (−0.22) = 295.72 K")
add_para("So cell 5 cools slightly (~0.03 K in 60 s) — consistent with the daily wave propagating from above.")

add_h("Step 4 — Apply bottom BC", level=2)
add_para("T_new[200] = T_new[199] (≈ 292.50 K both — bottom barely moves)")

add_h("Step 5 — Apply top BC", level=2)
add_para("T_new[0] = T_s_0_new = 299.9999 K. Replaces whatever the interior update tried to compute for cell 0 (if anything).")

add_h("Step 6 — Advance", level=2)
add_para("T = T_new. t = 60 s. Record T[0]=299.9999 K, G_surface = 0.30 · (299.9999 − T_new[1]) / 0.01 W/m². Loop back to Step 1.")

# -----------------------------------------------------------------
add_h("17. Why This Matters for Verification", level=1)
add_para(
    "After integrating for 5 days following the recipe above, we compare the day-5 cycle of T(z=10cm, t) "
    "and G(z=0, t) against the analytical damping-depth solution. The match should be excellent if:"
)
add_bullet("The IC was in quasi-equilibrium → no spin-up artefact remains by day 5.")
add_bullet("The harmonic-mean and flux divergence are correctly coded → no energy-conservation errors.")
add_bullet("The top BC is re-applied every step → the surface temperature tracks the prescribed sinusoid.")
add_bullet("The bottom BC zeroes the deep flux → no spurious leakage out the bottom.")
add_bullet("The chosen ν (for FTCS) is below 1/2 → no exponential noise growth.")
add_callout(
    "If any of these is broken, Test 1 fails. The whole point of Test 1 is to catch implementation "
    "bugs before we move on to Test 2 (whose results have no closed-form to check against)."
)

# -----------------------------------------------------------------
add_h("Appendix — The atmospheric forcing functions", level=1)
add_para("Synthetic, fully reproducible, sinusoidal:")
add_equation_block("S↓(t) = max[1000 · cos(ω(t − 12 h)), 0]  W/m²")
add_equation_block("L↓(t) = 350 + 20 · cos(ω(t − 14 h))     W/m²")
add_equation_block("T_a(t) = 292.5 + 7.5 · cos(ω(t − 14 h)) K")
add_equation_block("U(t) = 3 m/s   (constant)")
add_para("with ω = 2π / 86400 s ≈ 7.27 × 10⁻⁵ rad/s (one full cycle per day). Solar peaks at noon; air temperature and longwave peak 2 hours later at 14:00 LT, modelling the typical afternoon thermal lag.")

# -----------------------------------------------------------------
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— End of cheat-sheet —")
r.italic = True
r.font.color.rgb = GREY
r.font.size = Pt(10)

doc.save("CLIM715_Experiment_Cheatsheet.docx")
print("Wrote CLIM715_Experiment_Cheatsheet.docx")
